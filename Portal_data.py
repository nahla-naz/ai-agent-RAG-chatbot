import requests
from bs4 import BeautifulSoup
import urllib.parse
from dotenv import load_dotenv
import os
from docx import Document
from docx2pdf import convert



class Webscrape_pages():

    ''' Class to webscrape pages in portal '''

    def __init__(self):

            
                self.BASE_URL = ""  
                

                self.LOGIN_URL = urllib.parse.urljoin(self.BASE_URL, "")  

                self.HEADERS = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                                "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                    "Referer": self.BASE_URL,
                }

                self.pdf_path = "./DOCUMENTS/Webscraped_output.pdf"

                ## URL endpoints
                self.URL = [" ",
                " ",
                          
                    ]
                
                
                self.USERNAME = None
                self.PASSWORD = None

                self.session = None


    def Get_credentials(self):
         
        ''' Function to get credentials for authentication'''

        load_dotenv(dotenv_path='credentials.env')

        self.USERNAME = os.getenv("USERNAME")
        self.PASSWORD = os.getenv("PASSWORD")
        
        print(f"User: {self.USERNAME}")


    def Get_session(self):
         
            ''' Function to get session ID after authentication '''

            self.session = requests.Session()
            self.session.headers.update(self.HEADERS)


            resp = self.session.get(self.LOGIN_URL, verify=False, timeout=15)
            print("GET Login form:", resp.status_code)

            soup = BeautifulSoup(resp.text, "html.parser")


            login_form = soup.find("form")

            text = soup.get_text()

            print("Text:",text)

            print("Login form:",login_form)

            if not login_form:
                raise SystemExit("Login form not found ")


            action = login_form.get("action")

            session_url = f"{self.BASE_URL}/CBQ/"

            post_url = urllib.parse.urljoin(session_url, action)
            print("Form POST URL:", post_url)


            payload = {}

            for inp in login_form.find_all("input"):
                name = inp.get("name")
                if not name:
                    continue
                typ = inp.get("type", "").lower()

                
                if typ in ["Log-In", "button"]:
                    continue

                if typ == "hidden":
                    # Add hidden fields like CSRF tokens, RememberMe, etc.
                    payload[name] = inp.get("value", "")
                elif name == "UserName":
                    payload[name] = self.USERNAME
                elif name == "Password":
                    payload[name] = self.PASSWORD
                else:
                    # 
                    payload[name] = inp.get("value", "")

            print("Payload:", payload)


            print("Payload:", payload)


            login_resp = self.session.post(post_url, data=payload, headers={"Referer": self.LOGIN_URL}, verify=False, timeout=15)

            print("Log in response: ",login_resp)
            print("POSTed to:", post_url)
            print("Login response status:", login_resp.status_code)


            login_soup = BeautifulSoup(login_resp.text, "html.parser")
            error_phrases = [
                "invalid username or password",
                "the supplied credential is invalid."
            ]

            page_text = login_resp.text.lower()

            print("Page text:", page_text)

            if any(phrase in page_text for phrase in error_phrases):
                print(" Login failed: detected error message.")
            else:
                print(" Login successful. Session is authenticated.")


            protected_url = urllib.parse.urljoin(session_url, "")  ## EDIT 
            protected = self.session.get(protected_url, verify=False, timeout=15)

            print("Protected page status:", protected.status_code)
            print(protected.text[:1000])  

    


    def Start_webscraping(self):
            
            ''' Webscrape through all pages and store data in PDF'''
            
            ## Webscrape other pages
            # response = session.get("", verify=False)
            # print(response.status_code)
            # print(response.text)

            for i, url_ in enumerate(self.URL):

                print(f"Webscraping Page {i}: {url_}")
                 
                next_page_url = urllib.parse.urljoin(self.BASE_URL, url_)  
                response = self.session.get(next_page_url, verify=False, timeout=15)

                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    page_text = soup.prettify()
                    #print("\nText: ",page_text)


                    ## Add text to document ##
                    try:
                        doc = Document("webscrape_output.docx")
                    except:
                        doc = Document()

                    doc.add_heading(f"URL({url_}) :")

                    doc.add_paragraph(page_text)

                    doc.save("webscrape_output.docx")

                else:
                    print("Failed to access the page:", next_page_url)
                    print("Status code:", response.status_code)
            
            print("Saving document as pdf...")
            ## Save document as pdf ##
            convert("webscrape_output.docx", self.pdf_path)

            


if __name__ == "__main__":
     
     print("Starting session...")
     
     getdata = Webscrape_pages()

     getdata.Get_credentials()
     getdata.Get_session()
     getdata.Start_webscraping()
     