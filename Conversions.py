import docx 
import csv
import docx2pdf
import whisper
from docx import Document


class Convert_4_RAG():

    ''' Class to convert files to PDF for RAG pipeline'''

    def __init__(self):

                self.doc = docx.Document()
                self.filepath = ".\\conv"
                


    def csv_2_pdf(self,file_input : str):
            
            
            from docx2pdf import convert

            with open(self.filepath + "\\" + file_input , newline='') as f:
                csv_reader = csv.reader(f) 

                csv_headers = next(csv_reader)
                csv_cols = len(csv_headers)

                table = self.doc.add_table(rows=2, cols=csv_cols)
                hdr_cells = table.rows[0].cells

                for i in range(csv_cols):
                    hdr_cells[i].text = csv_headers[i]

                for row in csv_reader:
                    row_cells = table.add_row().cells
                    for i in range(csv_cols):
                        row_cells[i].text = row[i]

            self.doc.add_page_break()
            self.doc.save(self.filepath + "\\" + "Output_csv.docx")

            print("Docx file saved successfully.")

            ## CONVERT TO PDF
            ## Save document as pdf ##
            convert(self.filepath + "\\" + "Output_csv.docx", "Output_csv.pdf")


    def ppt_2_pdf(self,fileinput: str):
          
        from pptxtopdf import convert

        print("Converting PPT to PDF...")

        convert(self.filepath + "\\" + fileinput, self.filepath + "\\" + "Output_ppt.pdf")

        print("\nPPT converted successfully.")


    def video_2_pdf(self,fileinput : str):
            
            from docx2pdf import convert
          
            print("\nLoading model...")
            # Load the model

            list_models = whisper.available_models()
            print(list_models)

            model = whisper.load_model("large-v3")  # ['tiny.en', 'tiny', 'base.en', 'base', 'small.en', 'small', 
                                                    #'medium.en', 'medium', 'large-v1', 'large-v2', 'large-v3',
                                                    # 'large', 'large-v3-turbo', 'turbo']

            # Path to your local video file
            video_path = self.filepath + "\\" + fileinput

            print("Transcribing...")
            # Transcribe the video
            result = model.transcribe(video_path)

            # Print the transcription text
            print(result["text"])

            # Save transcription to a file
            #with open(self.filepath + "\\" + "Output_transcript.txt", "w", encoding="utf-8") as f:
            #    f.write(result["text"])

            
            doc = Document()

            doc.add_paragraph(result["text"])
            doc.save(self.filepath + "\\" + "Output_Transcript.docx")

            print("Saving Transcript as pdf...")
            ## Save document as pdf ##
            convert(self.filepath + "\\" + "Output_Transcript.docx", self.filepath + "\\" + "Output_Transcript.pdf")

            print("Transcript saved successfully. ")

if __name__ == "__main__":
      

      Convert_files = Convert_4_RAG()

      #Convert_files.csv_2_pdf("")
      #Convert_files.ppt_2_pdf("")
      Convert_files.video_2_pdf("")



                    
                    

